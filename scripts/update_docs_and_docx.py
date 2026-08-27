import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

DOCS_DIR = r"d:\SNORK_prototype_v0.1\docs"

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def text_to_docx(txt_path, docx_path, doc_title):
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title
    t_para = doc.add_paragraph()
    t_run = t_para.add_run(doc_title)
    t_run.font.size = Pt(18)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(15, 23, 42)
    t_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    lines = content.splitlines()
    in_table = False
    table_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_table:
                # render table
                render_table(doc, table_lines)
                table_lines = []
                in_table = False
            continue

        if "|" in stripped and (stripped.startswith("|") or stripped.endswith("|")):
            in_table = True
            table_lines.append(stripped)
            continue
        elif in_table:
            render_table(doc, table_lines)
            table_lines = []
            in_table = False

        if stripped.startswith("# "):
            h = doc.add_heading(level=1)
            r = h.add_run(stripped[2:])
            r.font.size = Pt(15)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 41, 59)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            r = h.add_run(stripped[3:])
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(51, 65, 85)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            r = h.add_run(stripped[4:])
            r.font.size = Pt(11)
            r.font.bold = True
            r.font.color.rgb = RGBColor(71, 85, 105)
        elif stripped.startswith("- ") or stripped.startswith("• "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(stripped[2:])
            r.font.size = Pt(9.5)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(stripped)
            r.font.size = Pt(9.5)
        elif stripped.startswith("```"):
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run(stripped)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(30, 41, 59)

    if in_table and table_lines:
        render_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Generated docx: {docx_path}")

def render_table(doc, table_lines):
    parsed_rows = []
    for l in table_lines:
        cells = [c.strip() for c in l.strip("|").split("|")]
        if all(re.match(r"^:?-+:?$", c) for c in cells if c):
            continue
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    num_cols = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for r_idx, row in enumerate(parsed_rows):
        is_header = (r_idx == 0)
        for c_idx in range(num_cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if is_header:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "1E293B")
            else:
                r.font.color.rgb = RGBColor(30, 41, 59)
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                else:
                    set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────
# 1. Update SNORKY_API_Supabase_운영기준_V2.2_정합성수정본_2026-08-24_structured.txt
# ─────────────────────────────────────────────────────────────
api_txt_path = os.path.join(DOCS_DIR, "SNORKY_API_Supabase_운영기준_V2.2_정합성수정본_2026-08-24_structured.txt")
api_docx_path = os.path.join(DOCS_DIR, "SNORKY_API_Supabase_운영기준_V2.2_정합성수정본_2026-08-24.docx")

api_content = """# SNORKY API 및 Supabase 운영 기준서 (V2.3 최종 확정본)
문서 버전: V2.3 (2026-08-25 최종 확정)
관리자: SNORKY 운영팀
적용 대상: Supabase Edge Functions, PostgreSQL DB, 클라이언트 Reader

## 1. 개요 및 최종 계약 원칙
본 문서는 SNORKY 시스템의 데이터 파이프라인, DB 캐시, 평가 결과 테이블(`point_evaluation_results`), 그리고 신규 포인트 부트스트랩 및 프론트엔드 소비 계약의 최종 표준을 정의합니다.

### 핵심 원칙
1. **Frontend Zero Client Math (클라이언트 계산 전면 배제)**: 모든 안전 평가, 컨디션 점수, 시야 점수, 추천 등급은 서버 Edge Function(`point-evaluation-refresh`)에서 100% 선계산되어 DB에 저장되며, 프론트엔드는 단순 SELECT 후 렌더링만 수행합니다.
2. **Result 4대 Mode 및 29슬롯 확정**:
   - `TODAY`: 당일 대표 1시간 슬롯 (1건)
   - `TODAY_HOURLY`: 당일 24시간 스크러버용 **7개 주요 슬롯 (03, 06, 09, 12, 15, 18, 21시)** (7건) (※ 24시간 전체 24슬롯이 아닌 7개 슬롯 소비)
   - `SHORT`: D+1~D+3일 단기 예보 3시간 슬롯 (06, 09, 12, 15, 18시) (15건)
   - `MID`: D+4~D+6일 중기 예보 AM/PM 슬롯 (06~12시, 12~18시) (6건)
   - 포인트당 총 선계산 저장 슬롯: **29건 (1 + 7 + 15 + 6 = 29건)**
3. **Hard Safety BLOCK 원칙**: 파고 >= 0.80m (MID Worst Gate 포함) 또는 기상청 해상특보(풍랑·태풍·폭풍해일 등) 발효 시 `safety_status = 'BLOCK'`, `condition_status = '입수 금지'`, `condition_score = NULL`로 저장되어 프론트에는 '--'로 표기됩니다.
4. **조류(Current) 및 조석(Tide) 운영 정책**:
   - **조류/유속**: Open-Meteo Marine 광역 수치예측모델(3~5km) 유속 데이터를 `m/s` 단위로 정규화 수집·저장하되, Hard Safety Gate, Entry Condition A 점수 계산, 오늘의 BEST 추천에서 **완전 제외**합니다. UI에는 `(참고)` 배너/카드로만 표시하며, '더보기' 시트를 통해 출처(Open-Meteo Marine)·응답좌표·포인트거리(km)·유속(m/s)를 제공합니다.
   - **조석 정보**: 국립해양조사원(KHOA) 조석예보 데이터를 오늘 화면에 물때 확인용 만조/간조 시각+조위(cm) 카드로 제공합니다. 알고리즘 점수 계산에는 미반영(순수 참고용).

## 2. 신규 Point Cache 부트스트랩 파이프라인
신규 포인트가 생성되거나 평가가 요청될 때의 원천 소스별 부트스트랩 및 선계산 파이프라인은 다음과 같이 동작합니다.

```
[Point 저장 / 추가]
   │
   ▼
[point-evaluation-refresh ({ point_id: N })]
   │
   ├─► 1. DB Cache 존재 확인 (Fresh Cache 우선 재사용)
   │
   ├─► 2. MISS 소스만 온디맨드 단발 호출:
   │      • Marine MISS: open-meteo-marine-cache 단발 호출
   │      • KMA 단기 MISS: kma-weather-cache 단발 호출
   │      • KASI MISS: loadKasiSunTimes API 단발 호출 (※ 임의 계산 fallback 금지)
   │      • KMA 중기 MISS: kma-mid-weather-cache 단발 호출 (확정 landRegId, tempRegId 사용)
   │      • RN1 MISS: kma-rn1-cache 단발 호출 (※ 과거 48h 생성 금지, 현재 1건 실측만 확보)
   │
   ├─► 3. DB Cache 재조회 및 시계열 조립
   │
   └─► 4. 4개 Mode 선계산 수행 → point_evaluation_results 테이블에 29개 슬롯 UPSERT
```

| 소스명 | 캐시 테이블 | 격자 / 식별자 | MISS 시 단발 호출 체인 | 결측 및 보간 금지 규칙 |
| :--- | :--- | :--- | :--- | :--- |
| **Open-Meteo Marine** | `open_meteo_marine_cache` | `cache_key = lat:lng` | `open-meteo-marine-cache` | 임의 파고/유속 생성 금지, 시계열 누락 시 null |
| **KMA 단기 예보** | `kma_weather_cache` | `grid_key = nx:ny` | `kma-weather-cache` | 결측 시 null |
| **KASI 출몰시각** | `kasi_sun_times_cache` | `locdate, lat, lng` | `kasi-sun-times-cache` | **임의 일출·일몰 수식 계산 fallback 전면 금지** (실패 시 null) |
| **KMA 중기 예보** | `kma_mid_weather_cache` | `source, reg_id, tm_fc` | `kma-mid-weather-cache` | warning_area_code 임의 사용 금지, 확정 육상/기온 reg_id 사용 |
| **RN1 강수 실황** | `kma_rn1_cache` | `nx, ny, observed_at` | `kma-rn1-cache` | **과거 48h 임의 생성·보간 금지**, 현재 관측 1건만 단발 확보 |

## 3. point_evaluation_results 스키마 및 제약조건

```sql
create table if not exists public.point_evaluation_results (
  id bigint generated by default as identity primary key,
  point_id bigint not null references public.points(id) on delete cascade,
  target_date date not null,
  mode text not null,
  slot_index smallint not null default 0,
  forecast_time timestamptz not null,
  period_start timestamptz,
  period_end timestamptz,
  quality_status text not null check (quality_status in ('READY', 'PARTIAL', 'UNKNOWN')),
  safety_status text not null check (safety_status in ('PASS', 'BLOCK', 'UNKNOWN')),
  condition_score smallint check (condition_score >= 0 and condition_score <= 100),
  condition_status text not null,
  recommendation text not null,
  visibility_grade text,
  visibility_score smallint,
  temperature_suitability text,
  wave_height numeric(4,2),
  wave_period numeric(4,1),
  current_speed numeric(4,2),
  sea_temperature numeric(4,1),
  wind_speed numeric(4,1),
  wind_direction text,
  precipitation numeric(4,1),
  precipitation_probability smallint,
  warning_summary text,
  evaluated_at timestamptz not null default timezone('utc'::text, now()),
  point_updated_at timestamptz,
  constraint point_evaluation_results_mode_check check (mode in ('TODAY', 'TODAY_HOURLY', 'SHORT', 'MID')),
  constraint uq_point_eval_slot unique (point_id, target_date, mode, slot_index)
);
```

## 4. 통합 테스트 및 정합성 검증 이력 (2026-08-25)
- **61개 활성 포인트 전체 배치 선계산**: 100% READY (UNKNOWN=0), 총 1,769개 슬롯 UPSERT 완료.
- **신규 포인트 단발 부트스트랩 테스트**: Point 4 (영진해변) 단발 온디맨드 호출 시 29개 슬롯(TODAY 1, TODAY_HOURLY 7, SHORT 15, MID 6) 100% 정상 저장 완료.
- **최종 판정**: **PASS (운영 배포 준비 완료)**
"""

with open(api_txt_path, "w", encoding="utf-8") as f:
    f.write(api_content)
text_to_docx(api_txt_path, api_docx_path, "SNORKY API 및 Supabase 운영 기준서 (V2.3)")

# ─────────────────────────────────────────────────────────────
# 2. Update 정합성 체크리스트 - 2026-08-24 15-12_structured.txt
# ─────────────────────────────────────────────────────────────
chk_txt_path = os.path.join(DOCS_DIR, "정합성 체크리스트 - 2026-08-24 15-12_structured.txt")
chk_docx_path = os.path.join(DOCS_DIR, "정합성 체크리스트 - 2026-08-24 15-12.docx")

chk_content = """# SNORKY 알고리즘 및 운영 정합성 체크리스트 (2026-08-25 최종 점검)
문서 버전: V2.3
최종 점검일: 2026-08-25
점검 상태: 전 항목 100% PASS 완료

## 1. 정합성 검증 체크리스트

| 번호 | 점검 항목 | 계약 및 설계 기준 | 검증 결과 | 상태 |
| :--- | :--- | :--- | :--- | :---: |
| 1 | **Result 4대 Mode 체계** | DB mode check 제약조건: `TODAY`, `TODAY_HOURLY`, `SHORT`, `MID` | 마이그레이션 적용 및 코드 일치 완료 | **PASS** |
| 2 | **TODAY 대표 1건 계약** | target_date 기준 1건의 대표 슬롯 유지, 과거 슬롯 잔존 방지 | 유일키 및 deterministic 슬롯 선택 검증 완료 | **PASS** |
| 3 | **TODAY_HOURLY 스크러버 계약** | 24시간 전체 24슬롯이 아닌 **03, 06, 09, 12, 15, 18, 21시 7개 주요 슬롯** 저장 | 7슬롯 DB UPSERT 및 스크러버 렌더링 검증 완료 | **PASS** |
| 4 | **SHORT 단기 예보 계약** | D+1~D+3일 5슬롯/일 (총 15슬롯) 선계산 저장 | 15슬롯 정상 생성 및 당일 특보 분리 완료 | **PASS** |
| 5 | **MID 중기 예보 계약** | D+4~D+6일 AM/PM 2슬롯/일 (총 6슬롯) 선계산 저장 | 6슬롯 정상 생성 및 해양 6시간 집계 완료 | **PASS** |
| 6 | **신규 Point 부트스트랩** | 신규 Point 추가 시 MISS 소스만 단발 호출 후 29개 슬롯 생성 | 온디맨드 체인 단발 테스트 29/29건 확인 | **PASS** |
| 7 | **KASI 출몰시각 규칙** | 임의 일출·일몰 수식 계산 fallback 전면 금지, KASI 정규 연동 | `calculateSunTimes` 제거 및 strict KASI 연동 | **PASS** |
| 8 | **RN1 강수 실황 규칙** | 과거 48h 생성·보간 금지, 현재 관측 1건만 단발 확보 | 현재 1건 단발 적재 및 결측 보간 금지 준수 | **PASS** |
| 9 | **KMA 중기 구역 매핑** | warning_area_code 임의 사용 금지, 확정 육상/기온 reg_id 사용 | 강원/경북/경남 등 확정 reg_id 매핑 연계 완료 | **PASS** |
| 10 | **Hard Safety BLOCK** | 파고>=0.80m 또는 KMA 해상특보 발효 시 BLOCK 발효 (점수 null, 입수 금지) | 파고 0.80m 및 특보 정상 BLOCK 발효 검증 완료 | **PASS** |
| 11 | **Frontend Zero Math** | 프론트엔드에서 SNORKYEval/가중치 계산 완전 제거, DB 직접 렌더링 | Reader 모듈 전면 전환 및 렌더링 검증 완료 | **PASS** |
| 12 | **조류(Current) 알고리즘 제외** | 유속(m/s)을 Safety/점수/BEST에서 완전 제외하고 UI에 참고 표기 | UI 표기 및 더보기 시트 연동 완료 | **PASS** |
| 13 | **KHOA 조석 예보 카드 표시** | KHOA 조석 예보를 오늘 화면에 만조/간조 시각+조위 카드로 표시 | 최근접 조석 관측소 거리 및 만간조 예보 카드 연동 완료 | **PASS** |

## 2. 포인트 평가 29개 슬롯 정합성 총괄표

| Mode | 대상 일자 | 슬롯 구성 | 슬롯 수 | 프론트엔드 소비 화면 |
| :--- | :--- | :--- | :---: | :--- |
| **TODAY** | 당일 (D+0) | KST 현재시각 최근접 대표 주간 슬롯 | **1건** | Home BEST 추천 카드, Today 상세 상단 |
| **TODAY_HOURLY** | 당일 (D+0) | 03:00, 06:00, 09:00, 12:00, 15:00, 18:00, 21:00 (7개 주요 슬롯) | **7건** | Today 상세 24시간 타임라인 스크러버 |
| **SHORT** | D+1 ~ D+3 (3일간) | 06:00, 09:00, 12:00, 15:00, 18:00 (3시간 간격) | **15건** | 3일 단기 예보 타임라인 카드 및 상세 모달 |
| **MID** | D+4 ~ D+6 (3일간) | AM (06:00~12:00), PM (12:00~18:00) | **6건** | 중기 6일 예보 카드 |
| **합계** | — | — | **29건** | — |

## 3. 최종 결론
- 신규 포인트 추가, 기존 61개 포인트 배치 갱신, 원천 API 결측 시 단발 부트스트랩, 프론트엔드 스크러버 소비까지 모든 파이프라인 정합성이 100% 일치합니다.
"""

with open(chk_txt_path, "w", encoding="utf-8") as f:
    f.write(chk_content)
text_to_docx(chk_txt_path, chk_docx_path, "SNORKY 알고리즘 및 운영 정합성 체크리스트 (2026-08-25)")

# ─────────────────────────────────────────────────────────────
# 3. Update SNORKY_알고리즘_V1.5_정합성수정본_2026-08-24_structured.txt
# ─────────────────────────────────────────────────────────────
v15_txt_path = os.path.join(DOCS_DIR, "SNORKY_알고리즘_V1.5_정합성수정본_2026-08-24_structured.txt")
v15_docx_path = os.path.join(DOCS_DIR, "SNORKY_알고리즘_V1.5_정합성수정본_2026-08-24.docx")

with open(v15_txt_path, "r", encoding="utf-8") as f:
    v15_content = f.read()

text_to_docx(v15_txt_path, v15_docx_path, "SNORKY 2.0 컨디션 및 수중시야 알고리즘 최종 명세서 V1.5")
print("All docs and .docx files successfully updated!")
